import re
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


SECTION_TITLES = {
    "单选": "一、选择题（单选）",
    "填空": "二、填空题",
    "作图": "三、作图题",
    "论述": "四、论述题",
    "实验": "五、实验探究题",
    "计算": "六、计算题",
}

TYPE_ORDER = ["单选", "填空", "作图", "论述", "实验", "计算"]

SCORES = {"单选": 2, "填空": 1, "作图": 2, "论述": 4, "实验": 1, "计算": 6}


def build_word(title: str, questions: list[dict], output_path: str):
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 标题
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 计算总分
    total = sum(SCORES.get(q["question_type"], 0) for q in questions)
    info = doc.add_paragraph(f"考试时间：90分钟    总分：{total}分")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 按题型分组输出
    grouped = {}
    for q in questions:
        grouped.setdefault(q["question_type"], []).append(q)

    num = 1
    for qtype in TYPE_ORDER:
        if qtype not in grouped:
            continue
        qs = grouped[qtype]
        score_each = SCORES.get(qtype, 0)
        total_score = score_each * len(qs)

        # 题型标题
        section_title = SECTION_TITLES.get(qtype, qtype)
        p = doc.add_paragraph()
        run = p.add_run(f"{section_title}（每题{score_each}分，共{total_score}分）")
        run.bold = True
        run.font.size = Pt(12)

        for q in qs:
            # 题目内容（去掉图片占位符，保留文字）
            content = q["content"].strip()
            content_text = re.sub(r'\[图:[^\]]+\]', '[图]', content)
            p = doc.add_paragraph(f"{num}. {content_text}")
            p.paragraph_format.space_after = Pt(4)
            num += 1

            # 选项逐行输出
            options = json.loads(q.get("options") or "[]")
            for opt in options:
                if opt.startswith("[图:"):
                    continue  # 选项图暂跳过
                opt_p = doc.add_paragraph(opt)
                opt_p.paragraph_format.left_indent = Cm(1)
                opt_p.paragraph_format.space_after = Pt(2)

            doc.add_paragraph()

    doc.save(output_path)


def build_pdf(word_path: str, pdf_path: str):
    import pythoncom
    from docx2pdf import convert
    pythoncom.CoInitialize()
    try:
        convert(word_path, pdf_path)
    finally:
        pythoncom.CoUninitialize()
